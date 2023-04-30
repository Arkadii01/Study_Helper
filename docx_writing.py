from bs4 import BeautifulSoup as BS
from flask import send_file
import requests
import docx
from docx.shared import Pt
import shutil
import os


def styling(doc):
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(14)

def parsing_save(name, links):
    for num in range(len(links)):
        responce = requests.get(links[num])
        html = BS(responce.text, 'html.parser')
        text = [i.text for i in html.findAll(['p', 'body > a', 'blockquote', 'h2', 'h1', 'h3', 'li', 'ol'])]
    
        doc = docx.Document()
        styling(doc)
        for par in text:
            doc.add_paragraph(par)
        
        doc.add_paragraph(f'Источник: {links[num]}')
        doc_name = f'{name}{num + 1}.docx'
        doc.save(doc_name)
        return send_file(f'/{doc_name}', as_attachment=True)


def summ_docx(docx_files):
    texts = []
    for docx_file in docx_files:
        doc = docx.Document(docx_file)
        styling(doc)

        text = []
        links = []
        for par in doc.paragraph:
            if 'Источник: ' in par:
                links.append(par[10:])
            else:
                text.append(par)
        texts.append('\n'.join(text))
    links = "\n".join(links)
    texts.append(f'Источники: {links}')

    doc = docx.Document()
    for par in texts:
        doc.add_paragraph(par)
    doc.save('Слияние.docx')